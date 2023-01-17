import requests
import json
from datetime import datetime, timedelta
from docx import Document, oxml, opc
from docx.enum.dml import MSO_THEME_COLOR_INDEX

# Set the time range for pull requests to be included (two weeks ago to now)
two_weeks_ago = datetime.now() - timedelta(weeks=2)
now = datetime.now()

def add_hyperlink(paragraph, url, text):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = oxml.shared.OxmlElement('w:r')
    rPr = oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def generate_table(doc, repo_owner, repo_name):
    # Make a GET request to the GitHub API to retrieve pull request information
    url = f'https://api.github.com/repos/{repo_owner}/{repo_name}/pulls?state=all&sort=updated&direction=desc&per_page=100'
    response = requests.get(url, headers={"Authorization": "Bearer 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"})
    data = json.loads(response.text)

    pull_requests = list(filter(
        (
            lambda pr: pr['state'] == 'open' or 
            (pr.get('closed_at', None) is not None and datetime.strptime(pr['closed_at'], '%Y-%m-%dT%H:%M:%SZ') > two_weeks_ago) or 
            (pr.get('merged_at', None) is not None and datetime.strptime(pr['merged_at'], '%Y-%m-%dT%H:%M:%SZ') > two_weeks_ago)
        ), 
        data))

    table = doc.add_table(rows=len(pull_requests)+1, cols=3)
    headers = table.rows[0].cells
    headers[0].text = "Pull Request"
    headers[1].text = "Comments"
    headers[2].text = "Status"

    # Iterate through the pull requests and add the relevant information to the document
    for i, pull_request in enumerate(pull_requests):
        pr_title: str = pull_request['title'].strip()
        if pr_title.endswith('.'):
            pr_title = pr_title[:-1]
        pr_title = "PR-{number}: {title}".format(number=pull_request['number'], title=pr_title)

        row = table.rows[i+1].cells
        paragraph = row[0].paragraphs[0]
        add_hyperlink(paragraph, pull_request['html_url'], pr_title)
        row[1].text = ""
        row[2].text = pull_request['state'].capitalize()




repos = [
    {"owner": "GPUOpen-LibrariesAndSDKs", "name": "RadeonProRenderMayaPlugin", "title": "Maya RPR"},
    {"owner": "GPUOpen-LibrariesAndSDKs", "name": "RadeonProRenderBlenderAddon", "title": "Blender RPR"},
    {"owner": "GPUOpen-LibrariesAndSDKs", "name": "BlenderUSDHydraAddon", "title": "Blender USD"},
    {"owner": "GPUOpen-LibrariesAndSDKs", "name": "RadeonProRenderUSD", "title": "Houdini"}
]


# Create a new Word document
doc = Document()

for repo in repos:
    doc.add_heading(repo['title'], level=1)
    generate_table(doc, repo['owner'], repo['name'])
    doc.add_page_break()

# Save the document
doc.save('pull_requests.docx')